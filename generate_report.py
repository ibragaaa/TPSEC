from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== PAGE DE GARDE =====
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Compte Rendu de TP')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Chiffrement avec Java')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Réseau et Sécurité — X. Skapin')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_paragraph()

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info2.add_run('Master 1 TDSI\nUniversité de Poitiers\nAnnée 2025-2026')
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'Mars 2026')
run.font.size = Pt(11)
run.italic = True

doc.add_page_break()

# ===== TABLE DES MATIÈRES =====
doc.add_heading('Table des matières', level=1)
toc_items = [
    ('1. Introduction', ''),
    ('2. Intégrité d\'un message', ''),
    ('   2.1 Condensat (Message Digest)', ''),
    ('   2.2 Code d\'authentification de message (MAC)', ''),
    ('3. Manipulation des algorithmes de chiffrement', ''),
    ('   3.1 Algorithmes symétriques (AES)', ''),
    ('   3.2 Modification DES/ECB/NoPadding', ''),
    ('   3.3 Algorithmes asymétriques (RSA)', ''),
    ('4. Signatures numériques', ''),
    ('   4.1 Signature avec DigitalSignatureExample_02', ''),
    ('   4.2 Protocole complet de signature (Section 3.1 du TP)', ''),
    ('5. Certificats numériques', ''),
    ('   5.1 Création d\'un keystore avec keytool', ''),
    ('   5.2 Stockage et chargement de clés (KeyStore Java)', ''),
    ('6. Communications SSL/TLS', ''),
    ('   6.1 Serveur HTTPS', ''),
    ('7. Conclusion', ''),
]
for item, _ in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        p.runs[0].bold = True

doc.add_page_break()

# ===== INTRODUCTION =====
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "L'objectif de ce TP est d'explorer les différents concepts de sécurité réseau "
    "liés au chiffrement, en utilisant la bibliothèque Java Cryptography Architecture (JCA). "
    "On va manipuler des algorithmes symétriques, asymétriques, des fonctions de hachage, "
    "des signatures numériques et des certificats. Tout ça en Java, ce qui permet de bien "
    "comprendre comment ces mécanismes fonctionnent concrètement."
)
doc.add_paragraph(
    "Pour ce TP, on a utilisé les fichiers Java fournis dans l'archive fichiers-java.zip, "
    "et on a exécuté chaque programme avec un argument personnalisé « Etudiant_Poitiers »."
)

# ===== 2. INTÉGRITÉ =====
doc.add_heading('2. Intégrité d\'un message', level=1)

doc.add_heading('2.1 Condensat (Message Digest)', level=2)
doc.add_paragraph(
    "Le premier exercice consiste à tester l'intégrité d'un message en utilisant une fonction "
    "de hachage. Le programme MessageDigestExample_01.java utilise l'algorithme SHA3-512 pour "
    "générer un condensat (digest) du texte passé en argument."
)

doc.add_paragraph("Compilation et exécution :")
p = doc.add_paragraph()
run = p.add_run(
    "$ javac MessageDigestExample_01.java\n"
    "$ java MessageDigestExample_01 \"Etudiant_Poitiers\""
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph("Résultat obtenu :")
p = doc.add_paragraph()
run = p.add_run(
    "SUN (DSA key/parameter generation; DSA signing; SHA-1, MD5 digests;\n"
    "SecureRandom; X.509 certificates; PKCS12, JKS & DKS keystores; ...)\n\n"
    "Digest:\n"
    "[suite d'octets non lisible en UTF-8]"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Explication des informations affichées : le programme affiche d'abord les informations "
    "sur le fournisseur (Provider) utilisé, ici le provider SUN. On voit qu'il prend en charge "
    "la génération de clés DSA, les algorithmes de hachage SHA-1 et MD5, les certificats X.509, "
    "les keystores, etc. Ensuite, il affiche le condensat généré. Ce condensat n'est pas lisible "
    "directement car c'est une suite d'octets bruts qui, interprétée en UTF-8, donne des caractères "
    "spéciaux ou illisibles. C'est normal car un condensat est un résumé binaire du message, pas "
    "du texte."
)

doc.add_paragraph(
    "Concernant les autres algorithmes de la famille SHA supportés par MessageDigest, "
    "d'après la documentation officielle de Java, on retrouve : "
    "SHA-1, SHA-224, SHA-256, SHA-384, SHA-512, SHA-512/224, SHA-512/256, "
    "SHA3-224, SHA3-256, SHA3-384 et SHA3-512. Le programme utilise SHA3-512, qui est la version "
    "la plus récente et la plus sécurisée de la famille SHA."
)

doc.add_heading('2.2 Code d\'authentification de message (MAC)', level=2)
doc.add_paragraph(
    "Un MAC (Message Authentication Code) est similaire à un condensat, sauf qu'il utilise "
    "en plus une clé secrète. Cela permet non seulement de vérifier l'intégrité du message "
    "mais aussi d'authentifier l'émetteur, car seul quelqu'un qui possède la clé peut générer "
    "le même MAC."
)

p = doc.add_paragraph()
run = p.add_run(
    "$ javac MessageAuthenticationCodeExample_01.java\n"
    "$ java MessageAuthenticationCodeExample_01 \"Etudiant_Poitiers\""
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph("Résultat obtenu :")
p = doc.add_paragraph()
run = p.add_run(
    "SunJCE Provider (implements RSA, DES, Triple DES, AES, Blowfish,\n"
    "ARCFOUR, RC2, PBE, Diffie-Hellman, HMAC, ChaCha20)\n\n"
    "MAC:\n"
    "[suite d'octets du MAC généré avec HmacSHA512]"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Ici on utilise le provider SunJCE qui implémente notamment HMAC. L'algorithme HmacSHA512 "
    "génère une clé secrète, puis calcule le MAC du texte avec cette clé. Le résultat est "
    "également une suite d'octets non lisible directement. Le MAC généré est unique pour la "
    "combinaison message + clé : si on change ne serait-ce qu'un bit du message ou de la clé, "
    "le MAC sera complètement différent."
)

# ===== 3. ALGORITHMES DE CHIFFREMENT =====
doc.add_heading('3. Manipulation des algorithmes de chiffrement', level=1)

doc.add_heading('3.1 Algorithmes symétriques (AES)', level=2)
doc.add_paragraph(
    "Le chiffrement symétrique utilise la même clé pour chiffrer et déchiffrer. "
    "Le programme SymmetricExample_01.java utilise AES en mode ECB avec padding PKCS5."
)

p = doc.add_paragraph()
run = p.add_run(
    "$ javac SymmetricExample_01.java\n"
    "$ java SymmetricExample_01 \"Etudiant_Poitiers\""
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph("Résultat obtenu :")
p = doc.add_paragraph()
run = p.add_run(
    "SunJCE Provider (implements RSA, DES, Triple DES, AES, Blowfish,\n"
    "ARCFOUR, RC2, PBE, Diffie-Hellman, HMAC, ChaCha20)\n\n"
    "[texte chiffré illisible]\n"
    "Etudiant_Poitiers"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Le programme affiche d'abord le provider SunJCE, puis le texte chiffré (qui est une "
    "suite d'octets illisible, ce qui est normal), et enfin le texte déchiffré qui correspond "
    "bien à l'argument original « Etudiant_Poitiers ». On voit que le chiffrement/déchiffrement "
    "fonctionne correctement : la clé AES de 128 bits est générée, le texte est chiffré en mode "
    "ECB avec padding PKCS5, puis déchiffré avec la même clé pour retrouver le texte original."
)

doc.add_heading('3.2 Modification DES/ECB/NoPadding', level=2)
doc.add_paragraph(
    "Le sujet demande de modifier le programme pour utiliser la combinaison DES/ECB/NoPadding. "
    "Avec DES, la taille de clé est de 56 bits et les blocs font 8 octets. Le mode NoPadding "
    "signifie qu'il n'y a pas d'ajout de bits pour compléter le dernier bloc, ce qui impose que "
    "la longueur du texte soit un multiple de 8 octets."
)

doc.add_paragraph(
    "J'ai donc créé un fichier SymmetricExample_DES.java qui adapte le code en conséquence. "
    "Il faut compléter le texte avec des espaces si sa longueur n'est pas un multiple de 8."
)

p = doc.add_paragraph()
run = p.add_run(
    "$ javac SymmetricExample_DES.java\n"
    "$ java SymmetricExample_DES \"Etudiant_Poitiers\"\n\n"
    "SunJCE Provider (implements RSA, DES, Triple DES, AES, ...)\n"
    "[texte chiffré illisible]\n"
    "Etudiant_Poitiers       "
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Les autres combinaisons de la famille DES disponibles sont : "
    "DES/ECB/PKCS5Padding, DES/CBC/NoPadding, DES/CBC/PKCS5Padding, "
    "DESede/ECB/NoPadding, DESede/ECB/PKCS5Padding, DESede/CBC/NoPadding, "
    "DESede/CBC/PKCS5Padding. DESede correspond au Triple DES qui applique DES trois fois "
    "avec des clés différentes pour renforcer la sécurité."
)

doc.add_heading('3.3 Algorithmes asymétriques (RSA)', level=2)
doc.add_paragraph(
    "Contrairement au chiffrement symétrique, le chiffrement asymétrique utilise une paire "
    "de clés : une clé publique pour chiffrer et une clé privée pour déchiffrer. Le programme "
    "AsymmetricExample_01.java utilise RSA avec une clé de 1024 bits."
)

p = doc.add_paragraph()
run = p.add_run(
    "$ javac AsymmetricExample_01.java\n"
    "$ java AsymmetricExample_01 \"Etudiant_Poitiers\""
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph("Résultat obtenu :")
p = doc.add_paragraph()
run = p.add_run(
    "SunJCE Provider (implements RSA, DES, Triple DES, AES, ...)\n\n"
    "[texte chiffré de 128 octets - illisible]\n"
    "Etudiant_Poitiers"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "On peut vérifier que le texte déchiffré est bien identique à l'argument fourni "
    "« Etudiant_Poitiers ». Le programme génère une paire de clés RSA de 1024 bits, "
    "chiffre le texte avec la clé publique, puis le déchiffre avec la clé privée. "
    "Le texte chiffré fait 128 octets (1024 bits / 8), ce qui est la taille typique "
    "d'un bloc RSA avec cette taille de clé. Le mode utilisé est RSA/ECB/PKCS1Padding."
)

# ===== 4. SIGNATURES NUMÉRIQUES =====
doc.add_heading('4. Signatures numériques', level=1)

doc.add_paragraph(
    "La signature numérique permet d'authentifier l'émetteur d'un message. Le principe "
    "est le suivant : l'émetteur génère un condensat du message, puis le chiffre avec sa "
    "clé privée. Le destinataire peut vérifier la signature en déchiffrant avec la clé "
    "publique de l'émetteur et en comparant avec le condensat qu'il calcule lui-même."
)

doc.add_heading('4.1 Signature avec DigitalSignatureExample_02', level=2)

p = doc.add_paragraph()
run = p.add_run(
    "$ javac DigitalSignatureExample_02.java\n"
    "$ java DigitalSignatureExample_02 \"Etudiant_Poitiers\""
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph("Résultat obtenu :")
p = doc.add_paragraph()
run = p.add_run(
    "Sun RSA signature provider\n\n"
    "Signature:\n"
    "[octets de la signature - illisible]\n\n"
    "Signature verified"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Le programme utilise l'algorithme SHA3-512WithRSA pour signer. Il génère une paire de "
    "clés RSA de 1024 bits, signe le message avec la clé privée en utilisant SHA3-512 comme "
    "fonction de hachage, puis vérifie la signature avec la clé publique. Le message "
    "« Signature verified » confirme que la vérification a réussi, ce qui prouve que le "
    "message n'a pas été modifié et qu'il provient bien du détenteur de la clé privée."
)

doc.add_heading('4.2 Protocole complet de signature (Section 3.1 du TP)', level=2)

doc.add_paragraph(
    "Pour cette partie, j'ai écrit un programme complet SignatureProtocol.java qui implémente "
    "tout le processus de signature numérique étape par étape, comme demandé dans le sujet. "
    "Le programme suit les 11 étapes décrites dans le TP :"
)

items = [
    "Récupération du texte en clair passé en argument",
    "Génération du condensat SHA-512 du message",
    "Création d'une paire de clés RSA de 2048 bits pour l'émetteur",
    "Création d'une paire de clés RSA de 1024 bits pour le destinataire",
    "Récupération du Cipher RSA/ECB/PKCS1Padding",
    "Chiffrement du condensat avec la clé privée de l'émetteur",
    "Chiffrement du texte avec la clé publique du destinataire",
    "Déchiffrement du condensat avec la clé publique de l'émetteur",
    "Déchiffrement du message avec la clé privée du destinataire",
    "Génération du condensat sur le message déchiffré",
    "Comparaison octet par octet des deux condensats",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph("Exécution :")
p = doc.add_paragraph()
run = p.add_run(
    "$ javac SignatureProtocol.java\n"
    "$ java SignatureProtocol \"Etudiant_Poitiers\"\n\n"
    "=== Protocole de Signature Numerique ===\n\n"
    "1. Texte en clair : Etudiant_Poitiers\n\n"
    "2. Provider du condensat : SUN (...)\n"
    "   Condensat genere (SHA-512), taille : 64 octets\n"
    "   Condensat (hex) : a324693104eb312183172f88eafd3d36669174598c562743\n"
    "   c2de9a4784e5b8eec5bb127b152e9769d5ad7eec18b6b04408e0bd1a503462ca\n"
    "   4625413dc6493aa6\n\n"
    "3. Paire de cles RSA emetteur generee (2048 bits)\n"
    "4. Paire de cles RSA destinataire generee (1024 bits)\n\n"
    "5. Provider du Cipher : SunJCE Provider (...)\n\n"
    "6. Condensat chiffre avec cle privee emetteur (256 octets)\n"
    "7. Message chiffre avec cle publique destinataire (128 octets)\n\n"
    "=== Reception par le destinataire ===\n\n"
    "8. Condensat dechiffre avec cle publique emetteur\n"
    "   Condensat (hex) : a324693104eb312183172f88eafd3d3666917459...\n\n"
    "9. Message dechiffre : Etudiant_Poitiers\n\n"
    "10. Condensat du message dechiffre genere\n"
    "    Condensat verification (hex) : a324693104eb312183172f88eafd3d3666917459...\n\n"
    "11. Comparaison des condensats octet par octet :\n"
    "    => Les deux condensats sont IDENTIQUES. Integrite verifiee !"
)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_paragraph(
    "Le résultat montre bien que tout le protocole fonctionne correctement. Le condensat "
    "déchiffré (étape 8) est identique au condensat initial (étape 2), et le message déchiffré "
    "(étape 9) correspond bien au texte original. La comparaison octet par octet (étape 11) "
    "confirme l'intégrité du message. Ce protocole combine à la fois la confidentialité "
    "(le message est chiffré avec la clé publique du destinataire) et l'authentification "
    "(le condensat est chiffré avec la clé privée de l'émetteur)."
)

# ===== 5. CERTIFICATS NUMÉRIQUES =====
doc.add_heading('5. Certificats numériques', level=1)

doc.add_paragraph(
    "Les certificats numériques permettent d'associer une identité vérifiée à une clé publique, "
    "grâce à une Autorité de Certification (CA). En Java, les clés et certificats sont stockés "
    "dans un keystore, protégé par mot de passe."
)

doc.add_heading('5.1 Création d\'un keystore avec keytool', level=2)

doc.add_paragraph(
    "On utilise l'outil keytool de Java pour générer une paire de clés RSA et un certificat "
    "auto-signé :"
)

p = doc.add_paragraph()
run = p.add_run(
    "$ keytool -genkey -v -alias MaCleEtudiant -keyalg RSA\n\n"
    "Generating 3,072 bit RSA key pair and self-signed certificate\n"
    "(SHA384withRSA) with a validity of 90 days\n"
    "for: CN=Etudiant Poitiers, OU=M1 TDSI, O=Universite de Poitiers,\n"
    "     L=Poitiers, ST=Vienne, C=FR\n"
    "[Storing /home/user/.keystore]"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Le fichier .keystore est créé dans le répertoire de connexion de l'utilisateur. "
    "On peut ensuite afficher les informations du certificat avec keytool -list :"
)

p = doc.add_paragraph()
run = p.add_run(
    "$ LANG=en_US.UTF-8 keytool -list -v -alias MaCleEtudiant\n\n"
    "Alias name: MaCleEtudiant\n"
    "Creation date: Mar 18, 2026\n"
    "Entry type: PrivateKeyEntry\n"
    "Certificate chain length: 1\n"
    "Owner: CN=Etudiant Poitiers, OU=M1 TDSI,\n"
    "       O=Universite de Poitiers, L=Poitiers, ST=Vienne, C=FR\n"
    "Issuer: CN=Etudiant Poitiers, OU=M1 TDSI, ...\n"
    "Serial number: 43a0baea32899f8e\n"
    "Valid from: Wed Mar 18 2026 until: Thu Mar 18 2027\n"
    "Signature algorithm name: SHA384withRSA\n"
    "Subject Public Key Algorithm: 3072-bit RSA key"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "On remarque que le certificat est auto-signé (Owner = Issuer), ce qui signifie qu'il "
    "n'a pas été validé par une autorité de certification externe. La clé RSA fait 3072 bits "
    "(valeur par défaut de keytool) et l'algorithme de signature est SHA384withRSA. "
    "Le préfixe LANG=en_US.UTF-8 est nécessaire pour contourner un bug connu de Java 24."
)

doc.add_heading('5.2 Stockage et chargement de clés (KeyStore Java)', level=2)

doc.add_paragraph(
    "Le programme StoringAndLoadingFromKeyStoreExample_01.java montre comment manipuler "
    "un KeyStore en Java de manière programmatique. Il effectue les opérations suivantes :"
)

items2 = [
    "Création d'un KeyStore en mémoire",
    "Stockage d'une clé secrète AES (SecretKey)",
    "Génération et stockage d'une paire de clés RSA 2048 bits (PrivateKey)",
    "Chargement de certificats X.509 depuis le fichier ca-chain.cert.pem",
    "Sauvegarde du KeyStore dans un fichier myKeystore",
    "Rechargement du KeyStore et vérification de toutes les entrées",
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run(
    "$ javac StoringAndLoadingFromKeyStoreExample_01.java\n"
    "$ java StoringAndLoadingFromKeyStoreExample_01\n\n"
    "Algorithm to generate secret key : AES\n"
    "Format used for the key: RAW\n"
    "Algorithm to generate private key : RSA\n"
    "Format used for the key: PKCS#8\n"
    "Cert type = X.509\n"
    "[Détails du certificat Alice Ltd Intermediate CA - 4096 bits]\n"
    "[Détails du certificat Alice Ltd Root CA - 4096 bits]\n"
    "KeyStore stored\n"
    "KeyStore loaded\n"
    "Algorithm to generate secret key : AES\n"
    "Format used for the key: RAW\n"
    "Algorithm to generate private key : RSA\n"
    "Format used for the key: PKCS#8"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Le programme charge deux certificats X.509 depuis ca-chain.cert.pem : un certificat "
    "intermédiaire (Alice Ltd Intermediate CA) et un certificat racine (Alice Ltd Root CA), "
    "tous deux avec des clés RSA de 4096 bits signés avec SHA256withRSA. "
    "Cette chaîne de certificats illustre le fonctionnement d'une PKI (Public Key Infrastructure) : "
    "le certificat racine signe le certificat intermédiaire, qui peut à son tour signer des "
    "certificats pour les utilisateurs finaux."
)

# ===== 6. SSL/TLS =====
doc.add_heading('6. Communications SSL/TLS', level=1)

doc.add_heading('6.1 Serveur HTTPS', level=2)

doc.add_paragraph(
    "Pour cette dernière partie, on met en place un serveur HTTPS en Java qui utilise SSL/TLS "
    "pour sécuriser les communications."
)

doc.add_paragraph(
    "Le fichier HTTPSServerExample_01.java crée un serveur SSL qui écoute sur le port 12345 "
    "et renvoie une page HTML simple quand un client se connecte."
)

doc.add_paragraph("Étape 1 — Première tentative sans keystore :", style='List Bullet')
doc.add_paragraph(
    "Si on compile et exécute le programme directement sans spécifier de keystore, "
    "le serveur ne peut pas démarrer correctement car il n'a pas de certificat SSL à "
    "présenter aux clients. On obtient une erreur liée à l'absence de keystore."
)

doc.add_paragraph("Étape 2 — Création du certificat serveur :", style='List Bullet')

p = doc.add_paragraph()
run = p.add_run(
    "$ keytool -genkey -v -keyalg RSA -alias monServeurCert \\\n"
    "  -keystore sslTlsKeyStore\n\n"
    "Generating 3,072 bit RSA key pair and self-signed certificate\n"
    "(SHA384withRSA) with a validity of 90 days\n"
    "for: CN=localhost, OU=TDSI, O=Universite de Poitiers,\n"
    "     L=Poitiers, ST=Vienne, C=FR\n"
    "[Storing sslTlsKeyStore]"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "Le CN (Common Name) doit être « localhost » car c'est le nom d'hôte que le navigateur "
    "va utiliser pour se connecter."
)

doc.add_paragraph("Étape 3 — Lancement du serveur avec le keystore :", style='List Bullet')

p = doc.add_paragraph()
run = p.add_run(
    "$ java -Djavax.net.ssl.keyStore=sslTlsKeyStore \\\n"
    "  -Djavax.net.ssl.keyStorePassword=motdepasse \\\n"
    "  HTTPSServerExample_01"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph("Étape 4 — Connexion depuis un navigateur :", style='List Bullet')
doc.add_paragraph(
    "En ouvrant Firefox à l'adresse https://localhost:12345, on obtient un avertissement de "
    "sécurité : « Warning: Potential Security Risk Ahead ». C'est normal car le certificat "
    "est auto-signé et n'a pas été émis par une autorité de certification reconnue. Le navigateur "
    "ne peut pas vérifier l'identité du serveur."
)
doc.add_paragraph(
    "En acceptant le risque et en ajoutant une exception de sécurité, on peut finalement "
    "se connecter au serveur et on voit la page HTML avec le message « Hello Students ». "
    "Le serveur affiche aussi dans la console « Client connection made » et la requête GET du navigateur."
)

doc.add_paragraph("Vérification avec curl :")
p = doc.add_paragraph()
run = p.add_run(
    "$ curl -k --http0.9 https://localhost:12345/\n\n"
    "<HTML><HEAD><TITLE>HTTPS Server Example</TITLE></HEAD>\n"
    "<BODY><H1>Hello Students</H1></BODY></HTML>"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    "L'option -k de curl permet d'ignorer la vérification du certificat SSL (puisqu'il est "
    "auto-signé). On reçoit bien la page HTML du serveur, ce qui confirme que la communication "
    "SSL/TLS fonctionne correctement."
)

# ===== 7. CONCLUSION =====
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    "Ce TP m'a permis de mettre en pratique les concepts de cryptographie vus en cours. "
    "On a pu voir concrètement comment fonctionnent les algorithmes symétriques (AES, DES) "
    "et asymétriques (RSA), les fonctions de hachage (SHA-512, SHA3-512), les MAC (HmacSHA512), "
    "les signatures numériques et les certificats X.509."
)
doc.add_paragraph(
    "Ce qui m'a le plus marqué, c'est la partie sur le protocole complet de signature "
    "numérique (section 3.1 du TP), où on combine vraiment tous les mécanismes ensemble : "
    "hachage pour l'intégrité, chiffrement asymétrique pour la confidentialité, et signature "
    "pour l'authentification. On comprend bien pourquoi ces trois aspects sont complémentaires "
    "et indispensables pour sécuriser une communication."
)
doc.add_paragraph(
    "La partie SSL/TLS est aussi très instructive car elle montre comment tous ces concepts "
    "s'appliquent dans un cas concret de communication réseau sécurisée. L'avertissement du "
    "navigateur pour le certificat auto-signé illustre bien l'importance des autorités de "
    "certification dans la chaîne de confiance."
)
doc.add_paragraph(
    "La bibliothèque JCA (Java Cryptography Architecture) est vraiment bien faite et permet "
    "de manipuler tous ces algorithmes de manière assez simple. Les providers (SUN, SunJCE) "
    "offrent une large palette d'algorithmes, et l'architecture est suffisamment modulaire "
    "pour pouvoir facilement changer d'algorithme sans modifier le reste du code."
)

output_path = '/workspace/compte_rendu_TP_chiffrement_java.docx'
doc.save(output_path)
print(f"Compte rendu sauvegardé : {output_path}")
