from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    doc.styles[f'Heading {i}'].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


# ===== PAGE DE GARDE =====
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Compte Rendu de TP')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Chiffrement avec Java')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Réseau et Sécurité — X. Skapin')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nicolas Moreau — Ibrahim Mejnoun')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Master 1 TDSI\nUniversité de Poitiers\nAnnée 2025-2026')
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mars 2026')
run.font.size = Pt(11)
run.italic = True

doc.add_page_break()

# ===== INTRO =====
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "Ce TP porte sur le chiffrement en Java. On va manipuler des algorithmes "
    "symétriques et asymétriques, des fonctions de hachage, des signatures numériques, "
    "des certificats et des communications SSL/TLS. Tout ça avec la bibliothèque JCA "
    "(Java Cryptography Architecture) qu'on a vue en cours."
)
doc.add_paragraph(
    "On a exécuté tous les programmes avec l'argument personnalisé « Etudiant_Poitiers ». "
    "On a aussi ajouté un affichage hexadécimal dans chaque programme pour que les résultats "
    "soient lisibles (les octets bruts en UTF-8 ne donnent rien d'exploitable dans le terminal)."
)

# ===== 2. INTÉGRITÉ =====
doc.add_heading('2. Intégrité d\'un message', level=1)

doc.add_heading('2.1 Condensat (Message Digest)', level=2)
doc.add_paragraph(
    "Comme vu en cours, pour tester l'intégrité d'un message on utilise une fonction de "
    "hachage qui produit un condensat (digest) : une suite de bits unique pour chaque message. "
    "Le programme MessageDigestExample_01.java utilise SHA3-512."
)

add_code(
    "$ javac MessageDigestExample_01.java\n"
    "$ java MessageDigestExample_01 \"Etudiant_Poitiers\""
)

doc.add_paragraph("Sortie obtenue :")
add_code(
    "SUN (DSA key/parameter generation; DSA signing; SHA-1, MD5 digests;\n"
    "SecureRandom; X.509 certificates; PKCS12, JKS & DKS keystores; PKIX\n"
    "CertPathValidator; PKIX CertPathBuilder; LDAP, Collection CertStores,\n"
    "JavaPolicy Policy; JavaLoginConfig Configuration)\n\n"
    "Digest (hex): 7295e9b04939f09111991a3a6cfdb0b9f4f00cac5567174c\n"
    "f17510405ec84396d6fec40c1ccf914097bc6fafe353094efd9fdd854\n"
    "8a3c6db5ac3f0ca6bd8bcc5"
)

doc.add_paragraph(
    "La première partie affiche le provider SUN, qui gère les algorithmes de hachage, "
    "les certificats X.509 et les keystores. Ensuite on a le condensat. À la base le "
    "programme l'affiche en UTF-8 et c'est complètement illisible (des caractères spéciaux "
    "partout), donc on a rajouté une ligne qui l'affiche en hex. "
    "Le condensat fait 512 bits (64 octets). Si on change ne serait-ce qu'un seul "
    "caractère du texte, le condensat change complètement — c'est le principe du hachage "
    "qu'on a vu en cours."
)

doc.add_paragraph(
    "Les autres algorithmes de la famille SHA supportés par MessageDigest sont : "
    "SHA-1, SHA-224, SHA-256, SHA-384, SHA-512, SHA-512/224, SHA-512/256, "
    "SHA3-224, SHA3-256, SHA3-384 et SHA3-512. En cours on a vu que MD5 et SHA-1 "
    "sont considérés comme obsolètes, et que SHA-2 et SHA-3 sont les standards actuels."
)

# --- 2.2 MAC ---
doc.add_heading('2.2 Code d\'authentification de message (MAC)', level=2)
doc.add_paragraph(
    "Un MAC c'est un condensat qui utilise en plus une clé secrète. Comme vu en cours, "
    "ça permet de vérifier à la fois l'intégrité du message et l'identité de l'émetteur, "
    "puisque seul quelqu'un qui a la clé peut produire le même MAC."
)
doc.add_paragraph(
    "On a édité le fichier MessageAuthenticationCodeExample_01.java pour comprendre ce "
    "qu'il fait : il génère une clé secrète HmacSHA512 avec KeyGenerator, puis calcule "
    "le MAC du texte avec un objet Mac."
)

add_code(
    "$ javac MessageAuthenticationCodeExample_01.java\n"
    "$ java MessageAuthenticationCodeExample_01 \"Etudiant_Poitiers\""
)

doc.add_paragraph("Sortie obtenue :")
add_code(
    "SunJCE Provider (implements RSA, DES, Triple DES, AES, Blowfish,\n"
    "ARCFOUR, RC2, PBE, Diffie-Hellman, HMAC, ChaCha20)\n\n"
    "MAC (hex): 510519e9a1d6a0a1ffdf6ae3f0ae3358a42f342680815bd4\n"
    "a63302045e2d6a3e83f80759d3d30f67a6ea9ac695003316c81c9b4b\n"
    "268cd6b26a061041ca8cf42a"
)
doc.add_paragraph(
    "Le provider ici c'est SunJCE. Le MAC fait 64 octets (512 bits). "
    "Comme la clé est générée aléatoirement à chaque exécution, le MAC change d'une "
    "fois sur l'autre. Mais pour un même message avec la même clé, on obtient toujours "
    "le même résultat."
)

# ===== 3. ALGORITHMES DE CHIFFREMENT =====
doc.add_heading('3. Manipulation des algorithmes de chiffrement', level=1)

doc.add_paragraph(
    "En cours on a vu que la qualité d'un algorithme de chiffrement dépend de la taille "
    "des blocs, du mode de chiffrement (comment les blocs sont enchaînés) et de la taille "
    "de la clé."
)

# --- 3.1 Symétrique ---
doc.add_heading('3.1 Algorithmes symétriques (AES)', level=2)
doc.add_paragraph(
    "Le chiffrement symétrique utilise la même clé pour chiffrer et déchiffrer — "
    "comme une serrure classique, comme expliqué en cours. "
    "SymmetricExample_01.java utilise AES en mode ECB avec remplissage PKCS5."
)

add_code(
    "$ javac SymmetricExample_01.java\n"
    "$ java SymmetricExample_01 \"Etudiant_Poitiers\""
)

doc.add_paragraph("Sortie :")
add_code(
    "SunJCE Provider (implements RSA, DES, Triple DES, AES, Blowfish,\n"
    "ARCFOUR, RC2, PBE, Diffie-Hellman, HMAC, ChaCha20)\n"
    "Chiffre (hex): 73ce0780523d948a5c597374c594f0381d0402a677a24fe6\n"
    "1b568a95341aa47b\n"
    "Etudiant_Poitiers"
)

doc.add_paragraph(
    "On voit le provider SunJCE, puis le texte chiffré en hex (32 octets, soit 2 blocs "
    "AES de 128 bits — le texte fait 17 caractères et le padding PKCS5 complète "
    "jusqu'à 32). La dernière ligne c'est le texte déchiffré, qui correspond bien "
    "à notre argument. Le chiffrement/déchiffrement fonctionne."
)

# --- 3.2 DES ---
doc.add_heading('3.2 Modification pour DES/ECB/NoPadding', level=2)
doc.add_paragraph(
    "On a modifié le programme pour utiliser DES/ECB/NoPadding. Avec DES la clé "
    "fait 56 bits et les blocs 8 octets. NoPadding veut dire qu'on ne rajoute pas "
    "de bits automatiquement, donc il faut que la longueur des données soit un multiple "
    "de 8. On a complété avec des espaces (en travaillant sur les bytes, pas les caractères, "
    "pour éviter les problèmes avec UTF-8)."
)

add_code(
    "$ javac SymmetricExample_DES.java\n"
    "$ java SymmetricExample_DES \"Etudiant_Poitiers\"\n\n"
    "SunJCE Provider (...)\n"
    "Chiffre (hex): 8ea8cc64b2c34929c18e370c00648abb4ad7ddae11064cdb\n"
    "Etudiant_Poitiers       "
)

doc.add_paragraph(
    "Le texte fait 17 octets, complété à 24 (3 blocs de 8). Le déchiffrement "
    "donne bien le texte original suivi des espaces de padding."
)

doc.add_paragraph(
    "Les autres combinaisons Algorithme/Mode/Remplissage de la famille DES "
    "disponibles dans Java :"
)
des_combos = [
    "DES/ECB/NoPadding (utilisé ici)",
    "DES/ECB/PKCS5Padding",
    "DES/CBC/NoPadding",
    "DES/CBC/PKCS5Padding",
    "DES/PCBC/NoPadding",
    "DES/PCBC/PKCS5Padding",
    "DES/CFB/NoPadding",
    "DES/CFB/PKCS5Padding",
    "DES/OFB/NoPadding",
    "DES/OFB/PKCS5Padding",
    "DES/CTR/NoPadding",
    "DES/CTS/NoPadding",
    "DESede/ECB/NoPadding",
    "DESede/ECB/PKCS5Padding",
    "DESede/CBC/NoPadding",
    "DESede/CBC/PKCS5Padding",
    "DESede/PCBC/NoPadding",
    "DESede/PCBC/PKCS5Padding",
    "DESede/CFB/NoPadding",
    "DESede/CFB/PKCS5Padding",
    "DESede/OFB/NoPadding",
    "DESede/OFB/PKCS5Padding",
    "DESede/CTR/NoPadding",
    "DESede/CTS/NoPadding",
]
for c in des_combos:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph(
    "DESede c'est le Triple DES (3DES), qui applique DES trois fois avec des clés "
    "différentes. ECB chiffre chaque bloc indépendamment, alors que CBC, CFB, OFB et "
    "PCBC utilisent un vecteur d'initialisation (IV) pour enchaîner les blocs entre eux. "
    "CTR transforme le chiffrement par blocs en chiffrement par flux, et CTS évite le "
    "padding. CTR et CTS ne marchent qu'avec NoPadding."
)

# --- 3.3 Asymétrique ---
doc.add_heading('3.3 Algorithmes asymétriques (RSA)', level=2)
doc.add_paragraph(
    "Le problème du symétrique c'est qu'il faut transmettre la clé au destinataire "
    "sans qu'elle soit interceptée. Les algorithmes asymétriques résolvent ça avec "
    "une paire de clés : une publique (pour chiffrer) et une privée (pour déchiffrer). "
    "C'est le principe qu'on a vu en cours avec RSA."
)
doc.add_paragraph(
    "En regardant le code d'AsymmetricExample_01.java, on identifie trois phases :"
)
doc.add_paragraph(
    "Génération de la paire de clés (lignes 16-18) : "
    "KeyPairGenerator.getInstance(\"RSA\") crée un générateur, initialisé à 1024 bits. "
    "generateKeyPair() produit la paire clé publique / clé privée.",
    style='List Bullet'
)
doc.add_paragraph(
    "Chiffrement (lignes 20-24) : "
    "le Cipher RSA/ECB/PKCS1Padding est initialisé en ENCRYPT_MODE avec la clé publique, "
    "puis doFinal(plainText) chiffre le texte.",
    style='List Bullet'
)
doc.add_paragraph(
    "Déchiffrement (lignes 26-29) : "
    "le même Cipher est réinitialisé en DECRYPT_MODE avec la clé privée, "
    "et doFinal(cipherText) déchiffre.",
    style='List Bullet'
)

add_code(
    "$ javac AsymmetricExample_01.java\n"
    "$ java AsymmetricExample_01 \"Etudiant_Poitiers\"\n\n"
    "SunJCE Provider (...)\n"
    "Chiffre (hex): 2e131577e3b6e81b3a8d731f4901049387901e5d\n"
    "9b12c6c3eec3dc053a028de070a6a6c6a2ea988b4a5fbcbe921da80f\n"
    "65da9f75d7b99bca7ea6972f4bac3a48f8e1627d49cf38b3e2d5aaad\n"
    "8f877c98e8f2c24461a9a97b4659321a9bb941cb2f14b1d1ce5f2a0f\n"
    "2d813cfd61ef4f3e62e8be29da75b85b3532a88870cd33da\n"
    "Etudiant_Poitiers"
)

doc.add_paragraph(
    "Le chiffré fait 128 octets (1024 bits / 8), ce qui correspond à la taille de clé RSA. "
    "Et le texte déchiffré est bien « Etudiant_Poitiers », identique à l'argument."
)

# ===== 4. SIGNATURES NUMÉRIQUES =====
doc.add_heading('4. Signatures numériques', level=1)

doc.add_paragraph(
    "En cours on a vu que la signature numérique permet d'authentifier l'émetteur. "
    "Le principe : l'émetteur génère un condensat du message, le chiffre avec sa clé "
    "privée, et envoie le tout. Le destinataire déchiffre le condensat avec la clé "
    "publique de l'émetteur et compare avec le condensat qu'il recalcule lui-même."
)

doc.add_heading('4.1 Vérification avec DigitalSignatureExample_02', level=2)

doc.add_paragraph(
    "Ce programme utilise SHA3-512WithRSA. Il génère une paire RSA de 1024 bits, "
    "signe le message avec la clé privée, puis vérifie la signature avec la clé publique."
)

add_code(
    "$ javac DigitalSignatureExample_02.java\n"
    "$ java DigitalSignatureExample_02 \"Etudiant_Poitiers\"\n\n"
    "Sun RSA signature provider\n\n"
    "Signature (hex): 4dba64ac0375ee688e14a943880ce6adc37e577d\n"
    "bbfaf30390ad9b78aa8c3dea8f25dd3355769e286142bb763494b9f2\n"
    "b3abc2d326a74f9c5dc371f1be353d524d8f57ffd448943b61ee70a3\n"
    "e06851e3aeb91b2b1e5873e0e14a20981700f1e89d10fdb42bca721b\n"
    "262b0233336c824d7cea67f0888066f963eb1df579b40d72\n\n"
    "Signature verified"
)

doc.add_paragraph(
    "« Signature verified » : la vérification a réussi, le message n'a pas été "
    "modifié et il vient bien du détenteur de la clé privée."
)

# --- 4.2 Protocole complet ---
doc.add_heading('4.2 Protocole complet de signature (Section 3.1 du TP)', level=2)

doc.add_paragraph(
    "On a écrit le programme SignatureProtocol.java qui implémente les 11 étapes "
    "demandées dans le sujet. Voici ce que fait chaque étape et le résultat obtenu."
)

doc.add_paragraph(
    "Étape 1 — On récupère le texte en byte[] depuis l'argument de la ligne de commande."
)
doc.add_paragraph(
    "Étape 2 — On génère le condensat avec SHA-512 (et pas SHA3-512 cette fois, "
    "comme demandé dans le sujet). Le provider est SUN."
)
add_code(
    "Condensat SHA-512 (64 octets) :\n"
    "a324693104eb312183172f88eafd3d36669174598c562743c2de9a47\n"
    "84e5b8eec5bb127b152e9769d5ad7eec18b6b04408e0bd1a503462ca\n"
    "4625413dc6493aa6"
)

doc.add_paragraph("Étape 3 — On génère une paire RSA de 2048 bits pour l'émetteur.")
doc.add_paragraph("Étape 4 — On génère une paire RSA de 1024 bits pour le destinataire.")
doc.add_paragraph(
    "Étape 5 — On récupère le Cipher RSA/ECB/PKCS1Padding (provider SunJCE). "
    "On utilise le même Cipher pour l'émetteur et le destinataire comme suggéré dans le sujet."
)
doc.add_paragraph(
    "Étape 6 — On chiffre le condensat avec la clé privée de l'émetteur. "
    "Sortie du programme :"
)
add_code(
    "6. Condensat chiffre avec cle privee emetteur (taille : 256 octets)\n"
    "   Condensat chiffre (hex) : 6605255d7f7ca7a5fe0d18a35210b8a8\n"
    "   f74e9727293b8edb7f5d316dcee9645a4817583c076a48423fad1477\n"
    "   86254a124e2b846079acfa043c63cf0e609e68139c6fd082871aed23\n"
    "   c3c92cc9b433321728d5c30c6a83a8c754eaacc9e0a6c41a07638e31\n"
    "   9f94f69f1a1af818395b27e5f34ee66cd999d9f6b75c6fe54cb3c25b\n"
    "   09f98f12d594f5c3ecf5fb6ce96869afe5a8d2f615cbc3c21be5840b\n"
    "   80ffe57a13ec6c4a2df079b0f4fe5aa3c43c10456be30ad20cf9ca07\n"
    "   bc765b8edf14fa45c0ffcd69360934d0107c7680c9c3a355df4003d4\n"
    "   75a7174c75cc820320aa5698e588c11b3cb71d5066c4e3baaad46db5\n"
    "   aa2fd6771e5db2f3666da84bda1d4f06"
)

doc.add_paragraph(
    "Étape 7 — On chiffre le texte en clair avec la clé publique du destinataire :"
)
add_code(
    "7. Message chiffre avec cle publique destinataire (taille : 128 octets)\n"
    "   Message chiffre (hex) : 2018f79ac9cfca8c5deabef007037f67\n"
    "   fdb3a55253953dc621bc446733b5617b870106e3a2ce69c77c1186a3\n"
    "   08d0069f60af5b051c59f714f9c3c9ff6816c6f561ecd7459e26dee0\n"
    "   22a0480484298d9310eadfcfabfc6d566f9b2b4633ca2e13b7b9add0\n"
    "   93cdf6e7e335ccf7b2b95b703ad5530ebb51f6143faf4dd72d7747e3"
)

doc.add_paragraph(
    "Étapes 8 à 11 — Côté destinataire, on déchiffre et on vérifie. "
    "Sortie du programme :"
)
add_code(
    "8. Condensat dechiffre avec cle publique emetteur\n"
    "   Condensat dechiffre (hex) : a324693104eb312183172f88eafd3d36\n"
    "   669174598c562743c2de9a4784e5b8eec5bb127b152e9769d5ad7eec\n"
    "   18b6b04408e0bd1a503462ca4625413dc6493aa6\n\n"
    "9. Message dechiffre : Etudiant_Poitiers\n\n"
    "10. Condensat du message dechiffre genere\n"
    "    Condensat verification (hex) : a324693104eb312183172f88eafd3d36\n"
    "    669174598c562743c2de9a4784e5b8eec5bb127b152e9769d5ad7eec\n"
    "    18b6b04408e0bd1a503462ca4625413dc6493aa6\n\n"
    "11. Comparaison des condensats octet par octet :\n"
    "    => Les deux condensats sont IDENTIQUES. Integrite verifiee !"
)

doc.add_paragraph(
    "Tout fonctionne : le condensat déchiffré (étape 8) est identique à celui "
    "de l'étape 2, le message déchiffré (étape 9) correspond à l'argument original, "
    "et la comparaison octet par octet confirme l'intégrité. "
    "Ce protocole combine la confidentialité (chiffrement avec la clé publique du "
    "destinataire) et l'authentification (chiffrement du condensat avec la clé "
    "privée de l'émetteur), comme on l'a vu en cours sur les signatures numériques."
)

# ===== 5. CERTIFICATS NUMÉRIQUES =====
doc.add_heading('5. Certificats numériques', level=1)

doc.add_paragraph(
    "Comme on l'a vu en cours, la signature numérique seule ne permet pas de savoir "
    "qui est vraiment l'émetteur. Les certificats numériques résolvent ce problème : "
    "ils contiennent une identité vérifiée et une clé publique, signées par une Autorité "
    "de Certification (CA) dans le cadre d'une PKI (Public Key Infrastructure). "
    "En Java, les clés et certificats sont stockés dans un keystore, géré avec keytool."
)

doc.add_heading('5.1 Création d\'un keystore avec keytool', level=2)

doc.add_paragraph(
    "On a utilisé keytool pour générer une paire de clés RSA avec un certificat auto-signé. "
    "On a choisi l'alias « MaCleEtudiant » :"
)
add_code(
    "$ keytool -genkey -v -alias MaCleEtudiant -keyalg RSA\n\n"
    "Generating 3,072 bit RSA key pair and self-signed certificate\n"
    "(SHA384withRSA) with a validity of 90 days\n"
    "for: CN=Etudiant Poitiers, OU=M1 TDSI,\n"
    "     O=Universite de Poitiers, L=Poitiers, ST=Vienne, C=FR\n"
    "[Storing /home/ubuntu/.keystore]"
)

doc.add_paragraph(
    "Le fichier créé est .keystore dans le répertoire de connexion. Le certificat est "
    "auto-signé (90 jours par défaut). La clé RSA fait 3072 bits."
)

doc.add_paragraph(
    "On affiche ensuite les informations de l'alias (le préfixe LANG=en_US.UTF-8 est "
    "nécessaire pour contourner un bug de Java, comme indiqué dans le sujet) :"
)
add_code(
    "$ LANG=en_US.UTF-8 keytool -list -v -alias MaCleEtudiant\n\n"
    "Alias name: MaCleEtudiant\n"
    "Creation date: Mar 18, 2026\n"
    "Entry type: PrivateKeyEntry\n"
    "Certificate chain length: 1\n"
    "Certificate[1]:\n"
    "Owner: CN=Etudiant Poitiers, OU=M1 TDSI,\n"
    "       O=Universite de Poitiers, L=Poitiers, ST=Vienne, C=FR\n"
    "Issuer: CN=Etudiant Poitiers, OU=M1 TDSI, ...\n"
    "Valid from: Wed Mar 18 07:58:27 UTC 2026\n"
    "      until: Tue Jun 16 07:58:27 UTC 2026\n"
    "Signature algorithm name: SHA384withRSA\n"
    "Subject Public Key Algorithm: 3072-bit RSA key"
)

doc.add_paragraph(
    "On voit que le certificat est auto-signé : Owner et Issuer sont identiques. "
    "Ça veut dire qu'il n'a pas été validé par une CA externe. "
    "La validité va du 18 mars au 16 juin 2026 (90 jours)."
)

# ===== 6. SSL/TLS =====
doc.add_heading('6. Communications SSL/TLS', level=1)

doc.add_paragraph(
    "SSL et TLS, qu'on a vus en cours, sont utilisés au niveau applicatif pour "
    "authentifier le serveur et le client. On va mettre en place un serveur HTTPS."
)

doc.add_heading('6.1 Analyse du code', level=2)
doc.add_paragraph(
    "Le fichier HTTPSServerExample_01.java crée un serveur HTTPS. "
    "Le code qui ouvre le serveur en écoute :"
)
add_code(
    "final int MYPORT = 12345;\n"
    "SSLServerSocketFactory sslsf =\n"
    "    (SSLServerSocketFactory)SSLServerSocketFactory.getDefault();\n"
    "ServerSocket ss = sslsf.createServerSocket(MYPORT);"
)
doc.add_paragraph(
    "SSLServerSocketFactory.getDefault() récupère la factory SSL par défaut, "
    "qui crée un ServerSocket SSL sur le port 12345 (supérieur à 1024, pas besoin "
    "de droits admin). Le reste du code attend un client (ss.accept()), lit sa requête, "
    "et renvoie une page HTML avec le message « Hello Students »."
)

doc.add_heading('6.2 Première exécution sans keystore', level=2)
add_code(
    "$ javac HTTPSServerExample_01.java\n"
    "$ java HTTPSServerExample_01"
)
doc.add_paragraph(
    "Le serveur démarre, mais quand on essaie de se connecter avec Firefox sur "
    "https://localhost:12345, le handshake SSL échoue. Côté serveur :"
)
add_code(
    "Client connection made\n"
    "javax.net.ssl.SSLHandshakeException:\n"
    "  (handshake_failure) No available authentication scheme"
)
doc.add_paragraph(
    "Le serveur n'a aucun certificat à présenter, donc il ne peut pas s'authentifier. "
    "Côté Firefox on voit « Secure Connection Failed » avec l'erreur "
    "SSL_ERROR_HANDSHAKE_FAILURE_ALERT."
)
doc.add_paragraph(
    "On arrête le programme avec Ctrl-C."
)

doc.add_heading('6.3 Création du certificat serveur', level=2)
doc.add_paragraph(
    "Pour résoudre le problème, on crée un certificat pour localhost dans un keystore "
    "sslTlsKeyStore. Le CN doit être « localhost » :"
)
add_code(
    "$ keytool -genkey -v -keyalg RSA -alias monServeurCert \\\n"
    "  -keystore sslTlsKeyStore\n\n"
    "Generating 3,072 bit RSA key pair and self-signed certificate\n"
    "(SHA384withRSA) with a validity of 90 days\n"
    "for: CN=localhost, OU=TDSI, O=Universite de Poitiers,\n"
    "     L=Poitiers, ST=Vienne, C=FR\n"
    "[Storing sslTlsKeyStore]"
)

doc.add_heading('6.4 Exécution avec le keystore', level=2)
doc.add_paragraph(
    "On vérifie que sslTlsKeyStore et HTTPSServerExample_01.class sont bien dans "
    "le même répertoire, puis on lance :"
)
add_code(
    "$ java -Djavax.net.ssl.keyStore=sslTlsKeyStore \\\n"
    "  -Djavax.net.ssl.keyStorePassword=motdepasse \\\n"
    "  HTTPSServerExample_01"
)
doc.add_paragraph("Le serveur démarre cette fois sans problème.")

doc.add_heading('6.5 Connexion avec Firefox', level=2)
doc.add_paragraph(
    "On ouvre Firefox à l'adresse https://localhost:12345. On obtient un avertissement "
    "de sécurité : « Warning: Potential Security Risk Ahead » avec le code "
    "SEC_ERROR_UNKNOWN_ISSUER. C'est normal : le certificat est auto-signé, il n'a pas "
    "été émis par une CA reconnue par Firefox."
)
doc.add_paragraph(
    "En cliquant sur « Advanced... » puis « Accept the Risk and Continue », on ajoute "
    "une exception. Firefox se connecte alors au serveur et affiche la page avec "
    "le message « Hello Students » en titre H1."
)
doc.add_paragraph("Côté serveur on voit :")
add_code(
    "Client connection made\n"
    "GET / HTTP/1.1"
)
doc.add_paragraph(
    "La connexion HTTPS fonctionne. Le serveur s'authentifie grâce au certificat "
    "dans le keystore. L'avertissement de Firefox vient du fait que le certificat "
    "n'est pas signé par une CA de confiance — en cours on a vu que c'est le rôle "
    "de la PKI de gérer cette chaîne de confiance."
)
doc.add_paragraph(
    "On peut aussi vérifier avec curl (l'option --http0.9 est nécessaire car le "
    "serveur renvoie du HTML brut sans ligne de statut HTTP) :"
)
add_code(
    "$ curl -k --http0.9 https://localhost:12345/\n\n"
    "<HTML><HEAD><TITLE>HTTPS Server Example</TITLE></HEAD>\n"
    "<BODY><H1>Hello Students</H1></BODY></HTML>"
)

# ===== 7. CONCLUSION =====
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    "Ce TP nous a permis de mettre en pratique les concepts de cryptographie vus en cours. "
    "On a manipulé les algorithmes symétriques (AES, DES), asymétriques (RSA), "
    "les fonctions de hachage (SHA-512, SHA3-512), les MAC (HmacSHA512), les signatures "
    "numériques et les certificats X.509."
)
doc.add_paragraph(
    "La partie la plus intéressante c'est le protocole complet de signature numérique "
    "(section 3.1), où on combine vraiment tout : le hachage pour l'intégrité, "
    "le chiffrement asymétrique pour la confidentialité, et la signature pour "
    "l'authentification. Ça rejoint directement ce qu'on a vu en cours sur les trois "
    "conditions de la signature (authentification, non-répudiation, non-falsification)."
)
doc.add_paragraph(
    "La partie SSL/TLS est aussi parlante parce qu'on voit concrètement comment "
    "le certificat auto-signé provoque un avertissement dans Firefox. "
    "Ça illustre bien le rôle des CA et de la PKI dans la chaîne de confiance."
)

output_path = '/workspace/TP_Chiffrement_Java/Compte_Rendu/compte_rendu_TP_chiffrement_java.docx'
doc.save(output_path)
print(f"Rapport sauvegardé : {output_path}")
